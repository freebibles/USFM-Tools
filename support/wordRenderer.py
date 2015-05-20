# -*- coding: utf-8 -*-
#

import abstractRenderer
import codecs
import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor

#
#   Renders to Word .docx
#

class WordRenderer(abstractRenderer.AbstractRenderer):
    
    def __init__(self, inputDir, outputFilename):
        self.document = Document()
        self.currentParagraph = None
        self.outputFilename = outputFilename
        self.outputFilename = outputFilename
        self.inputDir = inputDir
        self.f = codecs.open('/dev/null', 'w', 'utf_8_sig')
        self.currentChapter = ''
        self.currentBook = ''
        self.in_nd = False
        self.in_footnote = False
        # Old:
        # Flags
        self.printerState = {u'li': False, u'd': False, u'm': False}
        self.smallCapSections = True  # Sometimes we don't want to do this, like for Psalms
        self.justDidLORD = False
        self.justDidNB = False
        self.doNB = False
        self.narrower = False
        self.doChapterOrVerse = u''
        self.smallcaps = False

    def render(self, order='normal'):
        self.loadUSFM(self.inputDir)
        self.run(order)
        self.document.save(self.outputFilename)        
    
    #
    #   Support
    #

    def startNarrower(self, n):
        s = u'}' if self.narrower else u'\n\\blank[medium] '
        self.narrower = True
        s = s + u'\n\\noindentation \\Q{' + str(n) + u'}{'
        self.doNB = True
        return s

    def stopNarrower(self):
        s = u'}\n\\blank[medium] ' if self.narrower else u''
        self.narrower = False
        return s

    def escapeText(self, s):
        return s.replace('&', '\\&').replace('%', '\\%')
 
    def markForSmallCaps(self):
        if self.smallCapSections: 
             self.smallcaps = True

    def renderSmallCaps(self, s):
        if self.smallcaps == True:
            self.smallcaps = False
            return self.smallCapText(s)
        return s

    def smallCapText(self, s):
         i = 0
         while i < len(s):
             if i < 50:  #we are early, look for comma
                 if s[i] == u',' or s[i] == u';' or s[i] == u'(' or s[i:i+3] == u'and':
                     return u'{\sc ' + s[:i+1] + u'}' + s[i+1:]
             else: # look for space
                 if s[i] == ' ':
                     return u'{\sc ' + s[:i] + u'}' + s[i:]
             i = i + 1
         return u'{\sc ' + s + u'}'
         
    def startLI(self):
        if self.printerState[u'li'] == False:
            self.printerState[u'li'] = True
            #return u'\startitemize \item '
            return ur'\startexdent '
        else:
            #return u'\item '
            return ur'\par '
        
    def stopLI(self):
        if self.printerState[u'li'] == False:
            return u''
        else:
            self.printerState[u'li'] = False
            #return u'\stopitemize'
            return ur'\stopexdent '

    def startD(self):
        if self.printerState[u'd'] == False:
            self.printerState[u'd'] = True
        return u'\par {\startalignment[middle] \em '

    def stopD(self):
        if self.printerState[u'd'] == False:
            return u''
        else:
            self.printerState[u'd'] = False
            return u'\stopalignment }'
            
    def startM(self):
        r = self.stopD() + self.stopLI() + self.stopNarrower()
        self.printerState[u'm'] = True
        return r + ur'\par {\startalignment[flushleft] '

    def stopM(self):
        if self.printerState[u'm'] == False:
            return u''
        else:
            self.printerState[u'm'] = False
            return u'\stopalignment }'

    def newLine(self):
        s = u'\n\par \n'
        if self.doNB:
            self.doNB = False
            self.justDidNB = True
            s = s + ur'\noindentation '
        elif self.justDidNB:
            self.justDidNB = False
            s = s + ur'\indentation '
        return s
                    
    #
    #   Used helpers for Docx
    #
    def clean(self, token):
        return token.getValue().replace('~', ' ')
        
    def newPara(self):
        self.currentParagraph = self.document.add_paragraph()

    #
    #   Tokens
    #

    def render_h(self, token):       self.document.add_page_break(); self.currentBook = token.getValue()
    def render_mt(self, token):      self.document.add_heading(self.clean(token), level=0)
    def render_mt2(self, token):     self.document.add_heading(self.clean(token), level=1)
    def renderMS(self, token):       self.document.add_heading(self.clean(token), level=2)
    def renderMS2(self, token):      self.document.add_heading(self.clean(token), level=3)
    def render_p(self, token):       self.newPara(); 

    def render_s1(self, token):      self.document.add_heading(self.clean(token), level=4)
    def render_s2(self, token):      self.document.add_heading(self.clean(token), level=5)

    def renderC(self, token):        self.currentChapter = token.getValue()
    def renderV(self, token):
        if token.getValue() == '1':
            run = self.currentParagraph.add_run(self.currentBook + ' ' + self.currentChapter + ' ')
            run.font.color.rgb = RGBColor(255, 0, 0)
        else:
            run = self.currentParagraph.add_run(token.getValue() + ' ')
            run.font.color.rgb = RGBColor(0, 140, 0)
        run.font.superscript = True
            
    def renderTEXT(self, token):
        run = self.currentParagraph.add_run(self.clean(token) + ' ')
        if self.in_nd: 
            run.font.small_caps = True
        elif self.in_footnote:
            run.font.color.rgb = RGBColor(0, 0, 140)
        else:
            pass
        
    def render_q1(self, token):     self.newPara(); self.currentParagraph.paragraph_format.space_after = Pt(0)
    def render_q2(self, token):     self.newPara(); self.currentParagraph.paragraph_format.space_after = Pt(0); self.currentParagraph.add_run('\t')
    def render_q3(self, token):     self.newPara(); self.currentParagraph.paragraph_format.space_after = Pt(0); self.currentParagraph.add_run('\t\t')
    def renderNB(self, token):      self.newPara(); self.currentParagraph.left_indent = Inches(0)
    def renderB(self, token):       self.newPara(); self.currentParagraph.paragraph_format.space_after = Pt(0)

    def renderFS(self, token):
        run = self.currentParagraph.add_run(' [[ ')
        run.font.color.rgb = RGBColor(0, 0, 140)
        self.in_footnote = True
    def renderFE(self, token):
        run = self.currentParagraph.add_run(' ]] ')
        run.font.color.rgb = RGBColor(0, 0, 140)
        self.in_footnote = False
    
    def renderIS(self, token):      self.f.write( u'{\em ' )
    def renderIE(self, token):      self.f.write( u'} ' )
    def renderADDS(self, token):    self.f.write( u'{\em ' )
    def renderADDE(self, token):    self.f.write( u'} ' )

    def render_nd_s(self, token):   self.in_nd = True
    def render_nd_e(self, token):   self.in_nd = False
    def renderLI(self, token):      self.f.write( self.startLI() )
    def renderD(self, token):       self.f.write( self.startD() )
    def renderSP(self, token):      self.f.write( self.startD() )
    def renderPBR(self, token):     self.f.write( u' \\\\ ' )
    def renderFR(self, token):      self.f.write( u' ' + token.getValue() + u' ' )
    def renderFRE(self, token):     self.f.write( u' ' )
    def renderFK(self, token):      self.f.write( u' ' + token.getValue() + u' ' )
    def renderFT(self, token):      self.f.write( u' ' + token.getValue() + u' ' )
    def renderPI(self, token):      self.render_q(token)
    
    def render_is1(self, token):    self.render_s1(token)
    def render_ip(self, token):     self.render_p(token)
    def render_iot(self, token):    self.render_q1(token)
    def render_io1(self, token):    self.render_q2(token)
    
    def render_pb(self, token):     self.document.add_page_break()
    def render_m(self, token):      self.render_p(token)
    def render_periph(self, token): pass
    
    #
    #   Introductory codes
    #
    
    introTeXt = unicode(r""" """)
    closeTeXt = ur""" """